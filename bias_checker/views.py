from django.shortcuts import get_object_or_404, render, redirect
from django.urls import reverse
from docx import Document
from docx.shared import RGBColor
from django.conf import settings
import os
from .forms import FileForm
from .models import File
from bias_checker.static.bias_checker.dictionary import dictionary
import re
from django.contrib import messages
import xml.etree.ElementTree as ET
from docx.text.hyperlink import Hyperlink
from django.views.decorators.csrf import csrf_protect

from docx import Document
from docx.oxml import OxmlElement
from docx.oxml.ns import qn

def add_paragraph_bookmark_hyperlink(doc, bookmark_name, text):
    """
    Adds a paragraph that links to a bookmark within the document.
    
    :param doc: The document to which the paragraph will be added.
    :param bookmark_name: The name of the bookmark to link to.
    :param text: The text of the paragraph that will be a hyperlink.
    """

    # Add a new paragraph
    paragraph = doc.add_paragraph()
    
    # Create the hyperlink element pointing to the bookmark
    hyperlink = OxmlElement('w:hyperlink')
    hyperlink.set(qn('w:anchor'), bookmark_name)  # Internal link to bookmark
    
    # Create a run with the paragraph text
    run = OxmlElement('w:r')
    rPr = OxmlElement('w:rPr')

    # Add underline to the run properties
    underline = OxmlElement('w:u')
    underline.set(qn('w:val'), 'single')  # Single underline
    rPr.append(underline)
    run.append(rPr)

    text_element = OxmlElement('w:t')
    text_element.text = text
    run.append(text_element)
    hyperlink.append(run)
    
    # Append the hyperlink to the paragraph
    paragraph._element.append(hyperlink)

    # Add a newline after the hyperlink
    br = OxmlElement('w:br')  # Line break element
    paragraph._element.append(br)
    

def add_bookmark(paragraph, bookmark_name):
    """
    Adds a bookmark to a paragraph.
    
    :param paragraph: The paragraph where the bookmark will be added.
    :param bookmark_name: The name of the bookmark.
    """
    # Create a bookmark start element
    bookmark_start = OxmlElement('w:bookmarkStart')
    bookmark_start.set(qn('w:id'), '1')  # ID must be unique
    bookmark_start.set(qn('w:name'), bookmark_name)
    
    # Create a bookmark end element
    bookmark_end = OxmlElement('w:bookmarkEnd')
    bookmark_end.set(qn('w:id'), '1')
    
    # Add bookmark elements to the paragraph
    paragraph._element.insert(0, bookmark_start)  # Insert at the start of the paragraph
    paragraph._element.append(bookmark_end) 

def hasImage(par):
    """get all of the images in a paragraph 
    :param par: a paragraph object from docx
    :return: a list of r:embed 
    """
    ids = []
    root = ET.fromstring(par._p.xml)
    namespace = {
             'a':"http://schemas.openxmlformats.org/drawingml/2006/main", \
             'r':"http://schemas.openxmlformats.org/officeDocument/2006/relationships", \
             'wp':"http://schemas.openxmlformats.org/drawingml/2006/wordprocessingDrawing"}

    inlines = root.findall('.//wp:inline',namespace)
    for inline in inlines:
        imgs = inline.findall('.//a:blip', namespace)
        for img in imgs:     
            id = img.attrib['{{{0}}}embed'.format(namespace['r'])]
        ids.append(id)

    return ids

# Create your views here.

@csrf_protect
def index(request):
    try:    
        if request.method == "POST":
            form = FileForm(request.POST, request.FILES)
            context = {'form' : form}
            if form.is_valid():
                form.save()
            else:
                return render(request, "bias_checker/error.html", {"content":"FILE ALREADY EXISTS, Delete existing file to continue"})
            
            biasIdStart = 1
            genderBiasList = []
            genderBiasSentences= []
            gender = dictionary
            printingOnce = True
            last = File.objects.last()

            save_directory = os.path.join(settings.BASE_DIR, 'media')
            file_path = os.path.join(save_directory,str(last.file))#Get the saved file from submit
            doc = Document(file_path)
            # Create a folder to save images if it doesn't exist
            
            
            if doc.paragraphs != None:
                for p in doc.paragraphs:
                    original_text = p.text
                    parts = original_text.split(" ")
                    sentences = re.split(r"[,.]", original_text)
                    """
                        Logic to get the sentence where the gender-bias is located:
                            1. Separate the string using . and ,
                            2. check the sentence if a bias is a substring
                            3. save the sentence to the collection
                            4. separate the sentence using whitespaces
                            5. change the bias word to red font color
                    """
                    if hasImage(p):
                        print('image present')
                    else:
                        p.clear()
                        for part in parts:
                            #cleaned_data = re.sub(r"’s", '', part)
                            #cleaned_data = re.sub(r"[^A-Za-z0-9.-]", '', part)
                            styled_run = p.add_run(part + " ")
                            for word in gender:
                                isMatch = re.match(rf"\b{word}\b", part, re.IGNORECASE)
                                if isMatch:
                                    genderBiasList.append(word)
                                    styled_run.font.color.rgb = RGBColor(255,0,0)
                                    #print(len(genderBiasList))
                                    add_bookmark(p, f"{len(genderBiasList)}")
                                    #print(len(genderBiasList))

                    for sentence in sentences:
                        #print(sentence)
                        words = sentence.split(" ")
                        for word in words:
                            for bias in list(dict.fromkeys(genderBiasList)):
                                isMatch = re.match(rf"\b{bias}\b", word, re.IGNORECASE)
                                if isMatch:
                                    #print(bias, sentence)
                                    genderBiasSentences.append(sentence)
                                    if printingOnce:
                                        doc.add_paragraph("\n\nSentences that have Gender-bias words:\n")
                                    # Add another paragraph that links to the bookmark
                                    add_paragraph_bookmark_hyperlink(doc, f"{biasIdStart}", f"{sentence}")
                                    biasIdStart = biasIdStart + 1
                                    printingOnce = False
                                    #print(biasIdStart)

                    # for i in range(1, len(genderBiasList)+ 1):
                    #     for sent in genderBiasSentences:
                    #         print(i)
                    #         # Add another paragraph that links to the bookmark
                    #         add_paragraph_bookmark_hyperlink(doc, f"{i}", f"{sent}")

            if doc.tables != None:
                for table in doc.tables:
                    for row in table.rows:
                        for cell in row.cells:
                            for paragraph in cell.paragraphs:
                                original_text = paragraph.text
                                parts = original_text.split(" ")
                                sentences = re.split(r"[,.]", original_text)
                                if hasImage(paragraph):
                                    print('image present in cell')
                                else:
                                    paragraph.clear()
                                    for part in parts:
                                        #cleaned_data = re.sub(r"’s", '', part)
                                        #cleaned_data = re.sub(r"[^A-Za-z0-9.-]", '', part)
                                        styled_run = paragraph.add_run(part + " ")
                                        for word in gender:
                                            isMatch = re.match(rf"\b{word}\b", part, re.IGNORECASE)
                                            if isMatch:
                                                genderBiasList.append(word)
                                                styled_run.font.color.rgb = RGBColor(255,0,0)
                                                add_bookmark(paragraph, f"{len(genderBiasList)}")
                                
                                for sentence in sentences:
                                    words = sentence.split(" ")
                                    for word in words:
                                        for bias in list(dict.fromkeys(genderBiasList)):
                                            isMatch = re.match(rf"\b{bias}\b", word, re.IGNORECASE)
                                            if isMatch:
                                                #print(bias, sentence)
                                                genderBiasSentences.append(sentence)
                                                if printingOnce:
                                                    doc.add_paragraph("\n\nSentences that have Gender-bias words:\n")
                                                # Add another paragraph that links to the bookmark
                                                add_paragraph_bookmark_hyperlink(doc, f"{biasIdStart}", f"{sentence}")
                                                biasIdStart = biasIdStart + 1
                                                printingOnce = False

            #print(genderBiasSentences)

            if len(genderBiasList) <= 0:
                doc.add_paragraph("No gender-bias detected")
            else:
                doc.add_paragraph(f"{len(genderBiasList)} gender-bias detected")

            fileName = f"gender_checked_{last.id}_{last.file}"
            File.objects.filter(pk = last.id).update(file = fileName)# update the database to match the file
            doc.save(os.path.join(save_directory,fileName)) #save the file with a new filename
            os.remove(file_path) #remove the old file
            messages.add_message(request,messages.INFO, "File submitted. Click 'Files' button to view the output.")
        else:
            form = FileForm
            context = {'form' : form}
        return render(request, 'bias_checker/index.html', context)
    except:
        print("error in input")
        return render(request, 'bias_checker/error.html', {"content": "INTERNAL ERROR"})
    #return render(request, 'bias_checker/index.html', context)
    

def output(request):
    try:
        if request.method == "POST":
            file_id = request.POST.get("data-id")
            print(file_id)
            data = File.objects.filter(pk = file_id).first()
            print(data)
            save_directory = os.path.join(settings.BASE_DIR, 'media')
            file_path = os.path.join(save_directory, str(data.file))#Get the saved file from submit

            print(file_path)
            os.remove(file_path)
            data.delete()
            messages.add_message(request,messages.INFO, "File deleted")
            return redirect(reverse("bias_checker:output"))
        else:

            form = File.objects.all()
            context = {'form': form}
            #print(form)
            if not form.exists():
                print("Empty query")
                return render(request, "bias_checker/output.html", {"content": "No data"})

            return render(request, "bias_checker/output.html", context)
    except:
        print("error in output")
        return render(request, "bias_checker/error.html", {"content":"INTERNAL ERROR"})
    
def convert(request):
    try:
        if request.method == "POST":
            data:str = request.POST.get("hiddenTextArea")
            #cleaned_data = data.replace('</div>', "").replace('<span class="highlight-word">', "").replace('</span>',"")
            separated = data.split("<div>")
            print(separated)

            #create empty file
            File.objects.create(file = "generated_file.docx")
            last = File.objects.last()
            doc = Document()
            
            for row in separated:
                doc.add_paragraph(row)

            fileName = f"{last.id}_{last.file}"
            File.objects.filter(pk = last.id).update(file = fileName)# update the database to match the file
            save_directory = os.path.join(settings.BASE_DIR, 'media')
            file_path = os.path.join(save_directory, fileName)#Get the saved file from submit
            doc.save(file_path)
            #return redirect(reverse("bias_checker:index"))
            messages.add_message(request,messages.INFO, "File generated. Click 'Files' button to view the output.")
        return redirect(reverse("bias_checker:index"))
        return render(request, "bias_checker/index.html", {"success":"File generated is located at 'Files'"})
    except:
        return render(request, "bias_checker/error.html", {"content":"CONVERSION ERROR"})


# def error(request):
#     return render(request, "bias_checker/error.html", {"content":"Error in error"})